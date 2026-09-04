"""M4: check_output(골든 결과 검증)과 install --with-ui 램프 테스트."""

import importlib.util
import os
from pathlib import Path

import pytest

pytestmark = pytest.mark.deterministic
ROOT = Path(__file__).resolve().parents[1]
os.environ.setdefault("AI_WORK_SKILL_ROOT", str(ROOT))
os.environ.setdefault("DOCGEN_PROJECT_DIR", str(ROOT))
FIXTURES = ROOT / "mcp" / "docgen" / "fixtures"


def _load(rel: str):
    p = ROOT / rel
    spec = importlib.util.spec_from_file_location(p.stem.replace("-", "_"), p)
    m = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(m)
    return m


def test_check_output_pass_on_rendered(tmp_path):
    import sys

    sys.path.insert(0, str(ROOT / "mcp" / "docgen"))
    from docgen import core

    run = tmp_path / "run"
    run.mkdir()
    (run / "design.doc.md").write_text(
        (FIXTURES / "design.doc.md").read_text(encoding="utf-8"), encoding="utf-8"
    )
    (run / "gateway.deck.md").write_text(
        (FIXTURES / "gateway.deck.md").read_text(encoding="utf-8"), encoding="utf-8"
    )
    core.render_docx(str(run / "design.doc.md"), str(run / "design.docx"), base_dir=str(ROOT))
    core.render_pptx(str(run / "gateway.deck.md"), str(run / "gateway.pptx"), base_dir=str(ROOT))

    co = _load("eval/check_output.py")
    r = co.check_run(str(run))
    assert r["verdict"] == "PASS", r
    assert r["off_theme_colors"] == {} and not r["hard_violations"]


def test_check_output_flags_off_theme(tmp_path):
    from docx import Document
    from docx.shared import RGBColor

    d = Document()
    p = d.add_paragraph()
    run = p.add_run("off theme")
    run.font.color.rgb = RGBColor.from_string("FF00FF")  # 테마에 없는 색
    out = tmp_path / "bad.docx"
    d.save(str(out))

    co = _load("eval/check_output.py")
    colors = co.office_colors(out)
    assert "FF00FF" in colors


def test_theme_allowed_includes_derived_tint():
    co = _load("eval/check_output.py")
    allowed = co.theme_hexes(ROOT / "themes" / "datasolution.json")
    assert "E6F7FD" in allowed  # accent 10% 틴트(구성도 llm 노드)


def test_install_accent_ramp():
    inst = _load("templates/install.py")
    theme = {"colors": {"action": "#1366E3", "white": "#FFFFFF", "ink": "#222222"}}
    css = inst.accent_ramp(theme)
    for step in (100, 400, 700, 900):
        assert f"--ui-accent-{step}:" in css
    assert css.count("--ui-accent-") == 9
