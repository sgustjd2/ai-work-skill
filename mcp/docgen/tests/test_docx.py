import re
import zipfile

import pytest
from docx import Document

from docgen import docx_render, parse

pytestmark = pytest.mark.deterministic


@pytest.fixture
def rendered(design_md, theme, tmp_path):
    parsed = parse.parse_doc(design_md)
    out = tmp_path / "design.docx"
    res = docx_render.render_docx(parsed, str(out), theme, base_dir=str(tmp_path))
    return out, res


def _xml(path, member):
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        if member == "footer":
            m = next((n for n in names if re.match(r"word/footer\d+\.xml", n)), None)
            return z.read(m).decode("utf-8") if m else ""
        return z.read(member).decode("utf-8")


def test_reopens(rendered):
    out, _ = rendered
    d = Document(str(out))
    assert len(d.paragraphs) > 5
    assert len(d.tables) >= 3


def test_eastasia_font(rendered):
    out, _ = rendered
    assert 'w:eastAsia="맑은 고딕"' in _xml(out, "word/document.xml") or "w:eastAsia" in _xml(
        out, "word/styles.xml"
    )


def test_wordwrap_present(rendered):
    out, _ = rendered
    assert "w:wordWrap" in _xml(out, "word/styles.xml") or "w:wordWrap" in _xml(
        out, "word/document.xml"
    )


def test_page_number_field(rendered):
    out, _ = rendered
    foot = _xml(out, "footer")
    assert "PAGE" in foot and "NUMPAGES" in foot


def test_toc_and_update_fields(rendered):
    out, _ = rendered
    assert "TOC" in _xml(out, "word/document.xml")
    assert "updateFields" in _xml(out, "word/settings.xml")


def test_table_header_shading(rendered):
    out, _ = rendered
    assert re.search(r'w:shd[^>]*w:fill="[0-9A-Fa-f]{6}"', _xml(out, "word/document.xml"))


def test_caption_numbering(rendered):
    out, _ = rendered
    xml = _xml(out, "word/document.xml")
    # 구성도 캡션 → [그림 1], 표 캡션 → [표 1]
    assert "[그림 1]" in xml and "[표 1]" in xml


def test_marker_highlight(rendered):
    out, _ = rendered
    assert "yellow" in _xml(out, "word/document.xml").lower()


def test_diagram_embedded_as_image(rendered):
    out, _ = rendered
    with zipfile.ZipFile(str(out)) as z:
        media = [n for n in z.namelist() if n.startswith("word/media")]
    assert len(media) >= 1  # 구성도 PNG


def test_default_output_and_headings(rendered):
    _, res = rendered
    assert res["out_path"].endswith(".docx")
    assert len(res["headings"]) >= 5


def test_template_mode(theme, tmp_path):
    # 회사 양식 docx 를 만들고(자리표시자 + 사용자 스타일), 템플릿 모드로 렌더
    tpl = Document()
    tpl.styles["Normal"].font.name = "굴림"
    tpl.add_paragraph("{{title}}")
    tpl.add_paragraph("작성: {{author}}")
    tpl_path = tmp_path / "company.docx"
    tpl.save(str(tpl_path))

    parsed = parse.parse_doc("---\ntitle: 제목입니다\nauthor: 홍길동아님\n---\n# 1. 개요\n본문.\n")
    out = tmp_path / "from_template.docx"
    res = docx_render.render_docx(
        parsed, str(out), theme, base_dir=str(tmp_path), template_docx=str(tpl_path)
    )
    d = Document(str(out))
    texts = [p.text for p in d.paragraphs]
    assert any("제목입니다" in t for t in texts)  # 자리표시자 치환
    assert any("굴림" == (s.font.name or "") for s in d.styles if s.name == "Normal")  # 스타일 유지
    assert any(res["warnings"])
