"""블록 모델 → docx (FR-10). 표지·개정이력·목차·머리글/바닥글·스타일·표·구성도·캡션·마커.

색·폰트·치수는 테마에서만 온다(소스에 hex 리터럴 없음). 한글은 rFonts w:eastAsia 를 직접
지정해야 지정 폰트로 나온다.
"""

from __future__ import annotations

import re
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from . import theme as T
from .diagram import mermaid, png

_PLACEHOLDER = re.compile(r"\[[^\]]*(?:확인 필요|입력 필요|추가 필요|금액|일정)[^\]]*\]")
_TITLE_FIELD = re.compile(r"\{\{\s*(title|subtitle|date|author|version|org|security)\s*\}\}")


# ─── 저수준 헬퍼 ─────────────────────────────────────────────────────────────
def _set_rfonts(rpr, name: str) -> None:
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(a), name)


def _style_font(style, name, size_pt, color_hex=None, bold=False):
    style.font.name = name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    if color_hex:
        style.font.color.rgb = RGBColor.from_string(color_hex)
    _set_rfonts(style.element.get_or_add_rPr(), name)


def _ppr_keep_korean(style):
    """어절 단위 줄바꿈 관련 속성. [확인 필요] Word/PowerPoint 실측으로 확정(PRD D14)."""
    ppr = style.element.get_or_add_pPr()
    for tag, val in (("w:wordWrap", "1"), ("w:kinsoku", "1"), ("w:overflowPunct", "1")):
        el = ppr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            ppr.append(el)
        el.set(qn("w:val"), val)


def _run(
    paragraph,
    text,
    font_name,
    *,
    bold=False,
    italic=False,
    mono=None,
    color_hex=None,
    highlight=False,
    size_pt=None,
):
    run = paragraph.add_run(text)
    run.font.name = mono or font_name
    run.font.bold = bold
    run.font.italic = italic
    if size_pt:
        run.font.size = Pt(size_pt)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    _set_rfonts(run._element.get_or_add_rPr(), mono or font_name)
    return run


def _field(paragraph, instr):
    r = paragraph.add_run()
    fb = OxmlElement("w:fldChar")
    fb.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    fe = OxmlElement("w:fldChar")
    fe.set(qn("w:fldCharType"), "end")
    r._element.append(fb)
    r._element.append(it)
    r._element.append(fe)


def _shade(cell, hex_no_hash):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_no_hash)
    cell._element.get_or_add_tcPr().append(shd)


def _repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trpr.append(th)


def _rule(paragraph, hex_no_hash, size=18):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for a, v in (
        ("w:val", "single"),
        ("w:sz", str(size)),
        ("w:space", "1"),
        ("w:color", hex_no_hash),
    ):
        bottom.set(qn(a), v)
    pbdr.append(bottom)
    ppr.append(pbdr)


# ─── 스타일 정의 ─────────────────────────────────────────────────────────────
def _define_styles(doc, theme):
    ko = T.font(theme, "office_ko")
    mono = T.font(theme, "mono")
    d = theme["docx"]
    ink = T.rgb_hex(theme, "ink")
    head = T.rgb_hex(theme, d.get("heading_color", "primary"))

    normal = doc.styles["Normal"]
    _style_font(normal, ko, d["body_pt"], ink)
    normal.paragraph_format.line_spacing = d.get("line_spacing", 1.5)
    normal.paragraph_format.space_after = Pt(d.get("para_after_pt", 6))
    _ppr_keep_korean(normal)

    for name, size in (
        ("Heading 1", d["h1_pt"]),
        ("Heading 2", d["h2_pt"]),
        ("Heading 3", d["h3_pt"]),
    ):
        st = doc.styles[name]
        _style_font(st, ko, size, head, bold=True)
        st.paragraph_format.space_before = Pt(size * 0.9)
        st.paragraph_format.space_after = Pt(size * 0.35)
        st.paragraph_format.keep_with_next = True

    def ensure(name, base="Normal"):
        try:
            return doc.styles[name]
        except KeyError:
            return doc.styles.add_style(name, 1)  # WD_STYLE_TYPE.PARAGRAPH

    cap = ensure("Caption")
    _style_font(cap, ko, d.get("caption_pt", 9), T.rgb_hex(theme, "ink_muted"))
    code = ensure("CodeBlock")
    _style_font(code, mono, d.get("code_pt", 9), ink)
    quote = ensure("QuoteBlock")
    _style_font(quote, ko, d["body_pt"], T.rgb_hex(theme, "ink_muted"))
    tt = ensure("TableText")
    _style_font(tt, ko, d.get("table_pt", 9.5), ink)
    tt.paragraph_format.space_after = Pt(0)
    tt.paragraph_format.line_spacing = 1.2


# ─── 페이지·머리글·바닥글 ────────────────────────────────────────────────────
def _page_setup(doc, theme, fm):
    d = theme["docx"]
    sec = doc.sections[0]
    top, right, bottom, left = d.get("margins_mm", [25, 20, 25, 20])
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin, sec.right_margin = Mm(top), Mm(right)
    sec.bottom_margin, sec.left_margin = Mm(bottom), Mm(left)

    ko = T.font(theme, "office_ko")
    subtle = T.rgb_hex(theme, "ink_subtle")
    hp = sec.header.paragraphs[0]
    hp.paragraph_format.tab_stops.add_tab_stop(
        sec.page_width - sec.left_margin - sec.right_margin, WD_ALIGN_PARAGRAPH.RIGHT
    )
    _run(hp, fm.get("title", ""), ko, color_hex=subtle, size_pt=8)
    if fm.get("security"):
        _run(hp, "\t" + str(fm["security"]), ko, color_hex=subtle, size_pt=8)

    fpar = sec.footer.paragraphs[0]
    fpar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(fpar, "", ko, color_hex=subtle, size_pt=9)
    _field(fpar, "PAGE")
    _run(fpar, " / ", ko, color_hex=subtle, size_pt=9)
    _field(fpar, "NUMPAGES")


def _enable_update_fields(doc):
    el = doc.settings.element
    uf = OxmlElement("w:updateFields")
    uf.set(qn("w:val"), "true")
    el.append(uf)


# ─── 표지·개정이력·목차 ──────────────────────────────────────────────────────
def _cover(doc, theme, fm, logo_path=None):
    ko = T.font(theme, "office_ko")
    primary = T.rgb_hex(theme, "primary")
    if logo_path and Path(logo_path).exists():
        doc.add_picture(logo_path, width=Mm(35))
    rule_p = doc.add_paragraph()
    _rule(rule_p, primary, size=24)
    for _ in range(2):
        doc.add_paragraph()
    tp = doc.add_paragraph()
    _run(
        tp,
        fm.get("title", "제목 없음"),
        ko,
        bold=True,
        color_hex=primary,
        size_pt=theme["docx"]["h1_pt"] + 8,
    )
    if fm.get("subtitle"):
        sp = doc.add_paragraph()
        _run(
            sp,
            str(fm["subtitle"]),
            ko,
            color_hex=T.rgb_hex(theme, "ink_muted"),
            size_pt=theme["docx"]["h2_pt"],
        )
    for _ in range(3):
        doc.add_paragraph()
    meta = [
        ("조직", fm.get("org", "")),
        ("작성", fm.get("author", "")),
        ("일자", str(fm.get("date", ""))),
        ("버전", str(fm.get("version", ""))),
    ]
    tbl = doc.add_table(rows=0, cols=2)
    for k, v in meta:
        if not v:
            continue
        cells = tbl.add_row().cells
        _cell_text(cells[0], f"{k}", theme, bold=True)
        _cell_text(cells[1], f"{v}", theme)
    if fm.get("approvers"):
        doc.add_paragraph()
        ap = doc.add_paragraph()
        _run(
            ap,
            "결재",
            ko,
            bold=True,
            color_hex=T.rgb_hex(theme, "ink_muted"),
            size_pt=theme["docx"]["body_pt"],
        )
        appr = fm["approvers"] if isinstance(fm["approvers"], list) else [fm["approvers"]]
        t2 = doc.add_table(rows=2, cols=len(appr))
        t2.style = "Table Grid"
        for i, name in enumerate(appr):
            _cell_text(t2.rows[0].cells[i], str(name), theme, bold=True, center=True)
            _cell_text(t2.rows[1].cells[i], "", theme)
    doc.add_page_break()


def _history(doc, theme, fm):
    hist = fm.get("history")
    if not hist:
        return
    ko = T.font(theme, "office_ko")
    hp = doc.add_paragraph()
    _run(
        hp,
        "개정 이력",
        ko,
        bold=True,
        color_hex=T.rgb_hex(theme, "primary"),
        size_pt=theme["docx"]["h3_pt"],
    )
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    for i, h in enumerate(("버전", "일자", "작성자", "내용")):
        _cell_text(tbl.rows[0].cells[i], h, theme, bold=True, center=True)
        _shade(
            tbl.rows[0].cells[i],
            T.rgb_hex(theme, theme["docx"].get("table_header_fill", "surface")),
        )
    _repeat_header(tbl.rows[0])
    for row in hist:
        cells = tbl.add_row().cells
        vals = [
            row.get("version", ""),
            row.get("date", ""),
            row.get("author", ""),
            row.get("note", ""),
        ]
        for c, v in zip(cells, vals, strict=False):
            _cell_text(c, str(v), theme)
    doc.add_paragraph()


def _toc(doc, theme):
    ko = T.font(theme, "office_ko")
    p = doc.add_paragraph()
    _run(
        p,
        "목차",
        ko,
        bold=True,
        color_hex=T.rgb_hex(theme, "primary"),
        size_pt=theme["docx"]["h3_pt"],
    )
    tp = doc.add_paragraph()
    _field(tp, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_page_break()


# ─── 셀·표 ───────────────────────────────────────────────────────────────────
def _cell_text(cell, text, theme, *, bold=False, center=False, right=False, runs=None):
    cell.paragraphs[0].style = None
    p = cell.paragraphs[0]
    p.style = cell.part.document.styles["TableText"]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif right:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ko = T.font(theme, "office_ko")
    if runs:
        _emit_runs(p, runs, theme, base_bold=bold)
    else:
        _run(p, text, ko, bold=bold)


# ─── 본문 블록 ───────────────────────────────────────────────────────────────
def _emit_runs(paragraph, runs, theme, base_bold=False):
    ko = T.font(theme, "office_ko")
    mono = T.font(theme, "mono")
    for r in runs:
        text = r["text"]
        parts = _PLACEHOLDER.split(text)
        marks = _PLACEHOLDER.findall(text)
        for i, seg in enumerate(parts):
            if seg:
                _run(
                    paragraph,
                    seg,
                    ko,
                    bold=r["bold"] or base_bold,
                    italic=r["italic"],
                    mono=mono if r["code"] else None,
                )
            if i < len(marks):
                _run(paragraph, marks[i], ko, highlight=True)


class DocxRenderer:
    def __init__(self, theme, base_dir=None):
        self.theme = theme
        self.base_dir = Path(base_dir) if base_dir else Path.cwd()
        self.warnings: list[str] = []
        self._fig = 0
        self._tab = 0
        self._tmp: list[str] = []

    def render(self, parsed, out_path, logo_path=None):
        fm = parsed["frontmatter"]
        doc = Document()
        _define_styles(doc, self.theme)
        _page_setup(doc, self.theme, fm)
        _cover(doc, self.theme, fm, logo_path)
        _history(doc, self.theme, fm)
        if fm.get("toc", True):
            _toc(doc, self.theme)
            _enable_update_fields(doc)
        headings = []
        for b in parsed["blocks"]:
            self._block(doc, b, headings)
        self._footnotes(doc, parsed.get("footnotes", []))
        Path(out_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_path)
        for t in self._tmp:
            try:
                Path(t).unlink()
            except OSError:
                pass
        return {
            "out_path": str(Path(out_path).resolve()),
            "pages_estimate": max(1, len(parsed["blocks"]) // 6 + 2),
            "headings": headings,
            "warnings": self.warnings,
        }

    def _block(self, doc, b, headings):
        t = b["type"]
        if t == "heading":
            lvl = min(b["level"], 3)
            p = doc.add_paragraph(style=f"Heading {lvl}")
            _emit_runs(p, b["runs"], self.theme)
            headings.append({"level": b["level"], "text": b["text"]})
        elif t == "paragraph":
            p = doc.add_paragraph()
            _emit_runs(p, b["runs"], self.theme)
        elif t == "list":
            self._list(doc, b)
        elif t == "table":
            self._table(doc, b)
        elif t == "code":
            self._code(doc, b)
        elif t == "diagram":
            self._diagram(doc, b)
        elif t == "timeline":
            self._timeline(doc, b)
        elif t == "chart":
            self._chart(doc, b)
        elif t == "image":
            self._image(doc, b)
        elif t == "quote":
            for inner in b["blocks"]:
                if inner["type"] == "paragraph":
                    p = doc.add_paragraph(style="QuoteBlock")
                    _emit_runs(p, inner["runs"], self.theme)
        elif t == "pagebreak":
            doc.add_page_break()
        elif t == "hr":
            _rule(doc.add_paragraph(), T.rgb_hex(self.theme, "line"), size=6)

    def _list(self, doc, b):
        def emit(items, level):
            for it in items:
                style = "List Number" if b["ordered"] else "List Bullet"
                if level > 0:
                    style += f" {min(level + 1, 3)}"
                try:
                    p = doc.add_paragraph(style=style)
                except KeyError:
                    p = doc.add_paragraph(style="List Bullet")
                _emit_runs(p, it["runs"], self.theme)
                if it.get("children"):
                    emit(it["children"], level + 1)

        emit(b["items"], 0)

    def _table(self, doc, b):
        if b.get("caption"):
            self._tab += 1
            cap = doc.add_paragraph(style="Caption")
            _run(
                cap, f"[표 {self._tab}] {b['caption']}", T.font(self.theme, "office_ko"), bold=True
            )
        header, rows, align = b["header"], b["rows"], b["align"]
        ncol = len(header) or (len(rows[0]) if rows else 1)
        tbl = doc.add_table(rows=1, cols=ncol)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i in range(ncol):
            runs = header[i] if i < len(header) else []
            _cell_text(tbl.rows[0].cells[i], "", self.theme, bold=True, center=True, runs=runs)
            _shade(
                tbl.rows[0].cells[i],
                T.rgb_hex(self.theme, self.theme["docx"].get("table_header_fill", "surface")),
            )
        _repeat_header(tbl.rows[0])
        for row in rows:
            cells = tbl.add_row().cells
            for i in range(ncol):
                runs = row[i] if i < len(row) else []
                a = align[i] if i < len(align) else "left"
                _cell_text(
                    cells[i],
                    "",
                    self.theme,
                    right=(a == "right"),
                    center=(a == "center"),
                    runs=runs,
                )

    def _code(self, doc, b):
        p = doc.add_paragraph(style="CodeBlock")
        _shade_paragraph(p, T.rgb_hex(self.theme, "surface"))
        _run(p, b["text"], T.font(self.theme, "mono"), mono=T.font(self.theme, "mono"))

    def _image(self, doc, b):
        src = b["src"]
        path = Path(src)
        if not path.is_absolute():
            path = self.base_dir / src
        if path.exists():
            doc.add_picture(str(path), width=self._content_width(doc))
        else:
            self.warnings.append(f"그림 파일 없음: {src}")
        self._figure_caption(doc, b.get("caption"))

    def _diagram(self, doc, b):
        tmp = Path(tempfile.gettempdir()) / f"docgen_diag_{id(b)}.png"
        out, warns = png.to_png(b["spec"], self.theme, str(tmp), dpi=150)
        self.warnings += warns
        if out:
            self._tmp.append(out)
            doc.add_picture(out, width=self._content_width(doc))
        else:
            mm, _ = mermaid.to_mermaid(b["spec"], self.theme)
            p = doc.add_paragraph(style="CodeBlock")
            _run(
                p,
                "```mermaid\n" + mm + "\n```",
                T.font(self.theme, "mono"),
                mono=T.font(self.theme, "mono"),
            )
        self._figure_caption(doc, b.get("caption"))

    def _figure_caption(self, doc, text):
        if not text:
            return
        self._fig += 1
        cap = doc.add_paragraph(style="Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(cap, f"[그림 {self._fig}] {text}", T.font(self.theme, "office_ko"))

    def _timeline(self, doc, b):
        tasks = b["spec"].get("tasks", [])
        tbl_block = {
            "type": "table",
            "header": [[_r("항목")], [_r("시작")], [_r("종료")], [_r("비고")]],
            "rows": [
                [
                    [_r(str(t.get("label", "")))],
                    [_r(str(t.get("start", "")))],
                    [_r(str(t.get("end", "")))],
                    [_r(str(t.get("group", "")))],
                ]
                for t in tasks
            ],
            "align": ["left", "left", "left", "left"],
            "caption": b.get("caption") or b["spec"].get("title"),
        }
        self._table(doc, tbl_block)

    def _chart(self, doc, b):
        spec = b["spec"]
        cats = spec.get("categories", [])
        series = spec.get("series", [])
        header = [[_r(spec.get("title", "차트"))]] + [[_r(s.get("name", ""))] for s in series]
        rows = []
        for i, c in enumerate(cats):
            row = [[_r(str(c))]]
            for s in series:
                vals = s.get("values", [])
                row.append([_r(str(vals[i]) if i < len(vals) else "")])
            rows.append(row)
        unit = spec.get("unit", "")
        title = spec.get("title", "")
        cap = b.get("caption") or (f"{title} ({unit})" if unit else title)
        self._table(
            doc,
            {
                "type": "table",
                "header": header,
                "rows": rows,
                "align": ["left"] + ["right"] * len(series),
                "caption": cap,
            },
        )
        if spec.get("source"):
            sp = doc.add_paragraph(style="Caption")
            _run(sp, f"출처: {spec['source']}", T.font(self.theme, "office_ko"))

    def _footnotes(self, doc, footnotes):
        if not footnotes:
            return
        doc.add_paragraph()
        p = doc.add_paragraph(style="Heading 3")
        _run(
            p,
            "참고",
            T.font(self.theme, "office_ko"),
            bold=True,
            color_hex=T.rgb_hex(self.theme, "primary"),
        )
        for f in footnotes:
            fp = doc.add_paragraph(style="Caption")
            _run(fp, f"[{f['id']}] ", T.font(self.theme, "office_ko"), bold=True)
            _emit_runs(fp, f["runs"], self.theme)

    def _content_width(self, doc):
        sec = doc.sections[0]
        return sec.page_width - sec.left_margin - sec.right_margin


def _shade_paragraph(paragraph, hex_no_hash):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_no_hash)
    ppr.append(shd)


def _r(text):
    return {"text": text, "bold": False, "italic": False, "code": False}


# ─── 진입점 ──────────────────────────────────────────────────────────────────
def render_docx(parsed, out_path, theme, base_dir=None, logo_path=None, template_docx=None):
    if template_docx:
        return _render_template(parsed, out_path, theme, base_dir, template_docx)
    return DocxRenderer(theme, base_dir).render(parsed, out_path, logo_path)


def _render_template(parsed, out_path, theme, base_dir, template_docx):
    """회사 양식 docx 위에 렌더 — 스타일·머리글/바닥글·여백 유지, 본문 교체.
    # ponytail: 스타일 매핑은 layout_map.docx_styles(있으면), 없으면 우리 스타일 추가.
    """
    tpath = Path(template_docx)
    if not tpath.is_absolute() and base_dir:
        tpath = Path(base_dir) / template_docx
    doc = Document(str(tpath))
    fm = parsed["frontmatter"]
    body = doc.element.body

    # 자리표시자 치환(표지 등 유지), 값 채움
    subs = {
        k: str(fm.get(k, ""))
        for k in ("title", "subtitle", "date", "author", "version", "org", "security")
    }
    has_placeholder = _substitute_placeholders(doc, subs)

    # 본문 요소 제거(마지막 sectPr 유지). 표지에 자리표시자가 있었으면 표지도 유지됨.
    removed = 0
    if not has_placeholder:
        for el in list(body):
            if el.tag == qn("w:sectPr"):
                continue
            body.remove(el)
            removed += 1

    styles_avail = {s.name for s in doc.styles}
    if "Table Grid" not in styles_avail or "Heading 1" not in styles_avail:
        _define_styles(doc, theme)
    renderer = DocxRenderer(theme, base_dir)
    headings = []
    for b in parsed["blocks"]:
        renderer._block(doc, b, headings)
    renderer._footnotes(doc, parsed.get("footnotes", []))
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    warns = renderer.warnings + [f"템플릿 모드: 본문 요소 {removed}개 제거, 스타일·머리글 유지."]
    return {
        "out_path": str(Path(out_path).resolve()),
        "pages_estimate": max(1, len(parsed["blocks"]) // 6 + 1),
        "headings": headings,
        "warnings": warns,
    }


def _substitute_placeholders(doc, subs) -> bool:
    found = False
    for p in doc.paragraphs:
        text = "".join(r.text for r in p.runs)
        if _TITLE_FIELD.search(text):
            found = True
            new = _TITLE_FIELD.sub(lambda m: subs.get(m.group(1), ""), text)
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = new
    return found
